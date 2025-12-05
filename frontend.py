import streamlit as st
from ai_researcher_2 import INITIAL_PROMPT, graph, config
from pathlib import Path
import logging
import os
from langchain_core.messages import AIMessage, HumanMessage
import graphviz
from gtts import gTTS
import tempfile
from docx import Document
import io

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# --- Page Config ---
st.set_page_config(
    page_title="AI Researcher Pro",
    page_icon="🧬",
    layout="wide",
    initial_sidebar_state="expanded",
)

# --- Custom CSS (Glassmorphism & Modern UI) ---
st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;600;800&display=swap');

    html, body, [class*="css"] {
        font-family: 'Outfit', sans-serif;
    }

    /* Background Animation */
    .stApp {
        background: radial-gradient(circle at 10% 20%, rgb(33, 37, 41) 0%, rgb(15, 15, 15) 90%);
    }

    /* Glassmorphism Cards */
    .glass-card {
        background: rgba(255, 255, 255, 0.05);
        backdrop-filter: blur(10px);
        -webkit-backdrop-filter: blur(10px);
        border: 1px solid rgba(255, 255, 255, 0.1);
        border-radius: 16px;
        padding: 20px;
        margin-bottom: 20px;
        box-shadow: 0 4px 30px rgba(0, 0, 0, 0.1);
    }

    /* Gradient Text */
    .gradient-text {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        font-weight: 800;
        font-size: 3.5rem;
    }

    /* Chat Messages */
    .stChatMessage {
        background: rgba(255, 255, 255, 0.03);
        border: 1px solid rgba(255, 255, 255, 0.05);
        border-radius: 12px;
    }

    /* Sidebar */
    [data-testid="stSidebar"] {
        background: rgba(17, 24, 39, 0.95);
        border-right: 1px solid rgba(255, 255, 255, 0.05);
    }

    /* Buttons */
    .stButton button {
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
        color: white;
        border: none;
        border-radius: 8px;
        font-weight: 600;
        transition: transform 0.2s;
    }
    .stButton button:hover {
        transform: scale(1.02);
    }
    
    /* Quick Start Pills */
    .pill-container {
        display: flex;
        gap: 10px;
        flex-wrap: wrap;
        margin-bottom: 20px;
    }
    </style>
""", unsafe_allow_html=True)

# --- Session State Initialization ---
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
    st.session_state.chat_history.append({
        "role": "assistant", 
        "content": "👋 **Welcome to AI Researcher Pro!**\n\nI'm your autonomous research partner. I can:\n1. 🔍 **Scour Arxiv** for the latest papers.\n2. 📖 **Read & Analyze** complex PDFs.\n3. ✍️ **Write & Publish** professional research papers.\n\n*What are we discovering today?*"
    })

if "generated_papers" not in st.session_state:
    st.session_state.generated_papers = []  # List of dicts: {'path': str, 'topic': str, 'bookmarked': bool}

if "research_graph" not in st.session_state:
    st.session_state.research_graph = {"nodes": set(), "edges": set()}

if "bookmarks" not in st.session_state:
    st.session_state.bookmarks = []

# --- Helper Functions ---
def generate_audio(text):
    try:
        tts = gTTS(text=text, lang='en')
        with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as fp:
            tts.save(fp.name)
            return fp.name
    except Exception as e:
        logger.error(f"Audio generation failed: {e}")
        return None

def create_docx(text):
    doc = Document()
    doc.add_heading('AI Research Report', 0)
    doc.add_paragraph(text)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- Sidebar ---
with st.sidebar:
    st.markdown("## 🧬 Control Center")
    
    # Feature 5: Bookmarks System
    with st.expander("⭐ Bookmarks", expanded=False):
        if st.session_state.bookmarks:
            for i, bookmark in enumerate(st.session_state.bookmarks):
                st.markdown(f"**{i+1}. {bookmark['title']}**")
                if os.path.exists(bookmark['path']):
                    with open(bookmark['path'], "rb") as f:
                        st.download_button("⬇️", f, file_name=os.path.basename(bookmark['path']), key=f"bm_{i}")
        else:
            st.info("No bookmarks yet.")

    # Paper Archive
    if st.session_state.generated_papers:
        st.markdown("### 📂 Paper Archive")
        for i, paper in enumerate(st.session_state.generated_papers):
            file_path = paper['path']
            if os.path.exists(file_path):
                col_a, col_b = st.columns([0.8, 0.2])
                with col_a:
                    with open(file_path, "rb") as f:
                        st.download_button(
                            label=f"📄 {paper.get('topic', 'Paper')} ({i+1})",
                            data=f,
                            file_name=os.path.basename(file_path),
                            mime="application/pdf",
                            key=f"dl_{i}"
                        )
                with col_b:
                    # Bookmark Toggle
                    if st.button("⭐", key=f"star_{i}"):
                        st.session_state.bookmarks.append({'title': paper.get('topic', 'Paper'), 'path': file_path})
                        st.success("Saved!")
            else:
                st.warning(f"File not found: {os.path.basename(file_path)}")
        st.markdown("---")

    # Settings
    with st.expander("⚙️ Settings", expanded=True):
        model_type = st.selectbox("AI Model", ["Gemini 2.5 Pro", "Gemini 1.5 Flash (Fast)"])
        research_depth = st.select_slider("Depth", options=["Overview", "Standard", "Deep Dive"], value="Standard")
        
        # Feature 2: Multi-Language Support
        output_language = st.selectbox(
            "🌍 Output Language", 
            ["English", "Spanish", "French", "German", "Hindi", "Chinese", "Japanese"]
        )

        # Feature 1: AI Personas
        persona = st.selectbox(
            "🎭 Researcher Persona",
            ["The Professor (Academic)", "The Journalist (Engaging)", "The Skeptic (Critical)", "ELI5 (Simple)"]
        )

        # Feature 4: Critical Review Mode
        critical_mode = st.toggle("⚖️ Critical Review Mode")
    
    # Knowledge Graph Visualization
    if st.session_state.research_graph["nodes"]:
        st.markdown("### 🕸️ Knowledge Graph")
        graph = graphviz.Digraph()
        graph.attr(bgcolor='transparent')
        graph.attr('node', style='filled', fillcolor='#4F46E5', fontcolor='white', color='white')
        graph.attr('edge', color='white')
        
        for node in st.session_state.research_graph["nodes"]:
            graph.node(node)
        for edge in st.session_state.research_graph["edges"]:
            graph.edge(edge[0], edge[1])
            
        st.graphviz_chart(graph)

    # Clear Chat
    if st.button("🗑️ Reset Research Session", use_container_width=True):
        st.session_state.chat_history = []
        st.session_state.generated_papers = []
        st.session_state.research_graph = {"nodes": set(), "edges": set()}
        st.rerun()

# --- Main Content ---
col1, col2 = st.columns([0.7, 0.3])
with col1:
    st.markdown('<h1 class="gradient-text">TechSeva AI Researcher</h1>', unsafe_allow_html=True)

# Quick Start Topics (Only show if history is empty/short)
if len(st.session_state.chat_history) <= 1:
    st.markdown("### 🔥 Trending Research Topics")
    cols = st.columns(4)
    topics = [
        "🤖 Autonomous Agents", 
        "⚛️ Quantum Machine Learning", 
        "🧬 CRISPR Gene Editing", 
        "🌍 Climate Change AI"
    ]
    
    selected_topic = None
    for i, topic in enumerate(topics):
        if cols[i].button(topic, use_container_width=True):
            selected_topic = topic

    if selected_topic:
        user_input = f"I want to research {selected_topic}. Find me the latest papers."
    else:
        user_input = None
else:
    user_input = None

# Display Chat History
chat_container = st.container()
with chat_container:
    for msg in st.session_state.chat_history:
        avatar = "👤" if msg["role"] == "user" else "🧬"
        with st.chat_message(msg["role"], avatar=avatar):
            st.markdown(msg["content"])
            
            # Feature 3: Audio Summary
            if "audio_path" in msg:
                st.audio(msg["audio_path"])
            
            # Feature 2: Export to Word (for assistant messages)
            if msg["role"] == "assistant" and len(msg["content"]) > 100:
                docx_file = create_docx(msg["content"])
                st.download_button(
                    label="📝 Download as Word",
                    data=docx_file,
                    file_name="research_report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"docx_{msg['content'][:10]}"
                )

# Chat Input
if not user_input:
    user_input = st.chat_input("Ask me to research anything...")

if user_input:
    # Add user message
    st.session_state.chat_history.append({"role": "user", "content": user_input})
    with st.chat_message("user", avatar="👤"):
        st.markdown(user_input)

    # Prepare input for the agent
    # Inject Language, Persona, and Critical Mode Instructions
    system_prompt = INITIAL_PROMPT
    
    # Language
    if output_language != "English":
        system_prompt += f"\n\nIMPORTANT: Please write your final response and any research papers in {output_language}."
    
    # Persona
    if "Professor" in persona:
        system_prompt += "\n\nSTYLE: Adopt the persona of a distinguished Professor. Be academic, rigorous, and use formal language."
    elif "Journalist" in persona:
        system_prompt += "\n\nSTYLE: Adopt the persona of a Science Journalist. Be engaging, storytelling-driven, and accessible to the general public."
    elif "Skeptic" in persona:
        system_prompt += "\n\nSTYLE: Adopt the persona of a Skeptic. Question assumptions, look for methodological flaws, and be critical of claims."
    elif "ELI5" in persona:
        system_prompt += "\n\nSTYLE: Explain Like I'm 5. Use simple analogies, avoid jargon, and make it very easy to understand."

    # Critical Mode
    if critical_mode:
        system_prompt += "\n\nMODE: CRITICAL REVIEW. For every paper you analyze, provide a SWOT Analysis (Strengths, Weaknesses, Opportunities, Threats)."

    messages_for_agent = [{"role": "system", "content": system_prompt}]
    for msg in st.session_state.chat_history:
        if msg["role"] != "system":
             messages_for_agent.append(msg)

    chat_input = {"messages": messages_for_agent}
    
    logger.info("Starting agent processing...")

    # Stream response
    with st.chat_message("assistant", avatar="🧬"):
        message_placeholder = st.empty()
        full_response = ""
        
        with st.status("🧠 **Thinking & Researching...**", expanded=True) as status:
            try:
                from ai_researcher_2 import graph as agent_graph
                
                for event in agent_graph.stream(chat_input, config, stream_mode="values"):
                    message = event["messages"][-1]
                    
                    if getattr(message, "tool_calls", None):
                        for tool_call in message.tool_calls:
                            status.write(f"🔍 **Searching:** `{tool_call['name']}`")
                            logger.info(f"Tool call: {tool_call['name']}")
                            
                            if tool_call['name'] == 'arxiv_search':
                                topic = tool_call['args'].get('topic', 'Unknown Topic')
                                st.session_state.research_graph["nodes"].add("Research")
                                st.session_state.research_graph["nodes"].add(topic)
                                st.session_state.research_graph["edges"].add(("Research", topic))
                    
                    if message.type == "tool":
                        if message.name == "render_latex_pdf":
                            pdf_path = message.content
                            status.write(f"📄 **PDF Generated:** `{os.path.basename(pdf_path)}`")
                            if not any(p['path'] == pdf_path for p in st.session_state.generated_papers):
                                st.session_state.generated_papers.append({
                                    'path': pdf_path,
                                    'topic': "New Research Paper"
                                })
                        elif message.name == "arxiv_search":
                             status.write(f"✅ **Arxiv Results Received**")
                        else:
                            status.write(f"✅ **Data Received** from `{message.name}`")

                    if isinstance(message, AIMessage) and message.content:
                        if not message.tool_calls:
                            full_response = message.content
                            message_placeholder.markdown(full_response + "▌")
                
                status.update(label="**Research Complete!**", state="complete", expanded=False)
            
            except Exception as e:
                status.update(label="❌ **Error**", state="error")
                st.error(f"An error occurred: {str(e)}")
        
        # Final update
        message_placeholder.markdown(full_response)
        
        # Audio Summary
        audio_file = None
        if full_response and len(full_response) > 50:
            with st.spinner("🎧 Generating Audio Summary..."):
                audio_text = full_response[:500] + ("..." if len(full_response) > 500 else "")
                audio_file = generate_audio(audio_text)
                if audio_file:
                    st.audio(audio_file)
        
        # Feature 3: Smart Follow-up Questions (Simulated)
        # In a real app, we'd ask the LLM to generate these. For now, we'll suggest generic ones based on context or just static for demo.
        # Let's try to extract them from the response if possible, or just add generic ones.
        st.markdown("### 💡 Suggested Next Steps")
        c1, c2, c3 = st.columns(3)
        if c1.button("Tell me more about the methodology"):
            st.session_state.chat_history.append({"role": "user", "content": "Tell me more about the methodology used in these papers."})
            st.rerun()
        if c2.button("What are the limitations?"):
            st.session_state.chat_history.append({"role": "user", "content": "What are the main limitations or weaknesses of this research?"})
            st.rerun()
        if c3.button("Compare with other approaches"):
            st.session_state.chat_history.append({"role": "user", "content": "How does this compare to other state-of-the-art approaches?"})
            st.rerun()

        # Download PDF Button
        if st.session_state.generated_papers:
            latest_paper = st.session_state.generated_papers[-1]
            if os.path.exists(latest_paper['path']):
                st.success("🎉 **New Research Paper Ready!**")
                with open(latest_paper['path'], "rb") as f:
                    st.download_button(
                        label="⬇️ Download PDF Now",
                        data=f,
                        file_name=os.path.basename(latest_paper['path']),
                        mime="application/pdf",
                        key="dl_latest_main"
                    )
    
    # Save to history
    msg_data = {"role": "assistant", "content": full_response}
    if audio_file:
        msg_data["audio_path"] = audio_file
    
    st.session_state.chat_history.append(msg_data)
    # We don't rerun here to avoid losing the "Smart Follow-up" buttons immediately
    # But if we don't rerun, the history won't update until next interaction.
    # Streamlit's chat_message handles display immediately, so it's fine.
